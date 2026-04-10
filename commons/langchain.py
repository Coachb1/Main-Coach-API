import os 
import re
import openai
import requests
import anthropic
import time

from pytube.exceptions import VideoUnavailable
from urllib.parse import urlparse, parse_qs
from moviepy.editor import *
from pytube import YouTube

from langchain.chains.summarize import load_summarize_chain
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.document_loaders import TextLoader
from langchain.docstore.document import Document
from langchain_community.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain_community.llms import OpenAI
from datetime import datetime
from pydub import AudioSegment
from django.conf import settings
from commons.timeit import timeit



# Validation Functions
def is_valid_openai_key(api_key) -> bool:
    """
    Check if the provided OpenAI API key is valid.

    Parameters:
    - api_key (str): The OpenAI API key to be tested.

    Returns:
    - bool: True if the API key is valid, False otherwise.

    Raises:
    - requests.exceptions.HTTPError: If there is an HTTP error during the authentication request.
    - Exception: If there is an error during the OpenAI Python package test.

    Example:
    >>> is_valid_openai_key("your_api_key")
    True
    """
    try:
        # Test OpenAI API key by making a request to the authentication endpoint
        auth_response = requests.get("https://api.openai.com/v1/engines", headers={"Authorization": f"Bearer {api_key}"})
        auth_response.raise_for_status()

        # Test OpenAI Python package by creating an instance of the openai.api object
        openai.api_key = api_key
        openai.Completion.create(engine="text-davinci-002", prompt="Hello, World!")

        # If both tests pass, the API key is valid
        return True

    except (requests.exceptions.HTTPError, Exception):
        # If either test fails, the API key is invalid
        return False
    
def is_valid_youtube_url(url: str) -> bool:
    """
    Check if the given URL is a valid YouTube video URL.

    Parameters:
        url (str): The URL to be checked.

    Returns:
        bool: True if the video is available, False otherwise.
    """
    # Check if the URL is a valid YouTube video URL
    try:
        # Create a YouTube object
        yt = YouTube(url)

        # Check if the video is available
        if not yt.video_id:
            return False

    except (VideoUnavailable, Exception):
        return False

    # Return True if the video is available
    return yt.streams.filter(adaptive=True).first() is not None

# Calculate YouTube video duration
def get_video_duration(url: str) -> float:
    """
    Get the duration of a video from a given URL.

    Parameters:
    - url (str): The URL of the video.

    Returns:
    - float: The duration of the video in minutes.

    Example:
    >>> get_video_duration('https://www.youtube.com/watch?v=dQw4w9WgXcQ')
    3.5
    """
    yt = YouTube(url)  
    video_length = round(yt.length / 60, 2)

    return video_length

# Calculate API call cost
def calculate_api_cost(video_length: float, option: str) -> float:
    """
    Calculate the cost of API calls based on the length of the video and the option chosen.

    Parameters:
    video_length (float): The length of the video in seconds.
    option (str): The option chosen, either 'summary' or 'answer'.

    Returns:
    float: The cost of the API calls.
    """
    if option == 'summary':
        api_call_cost = round(video_length * 0.009, 2)
    elif option == 'answer':
        api_call_cost = round(video_length * 0.006, 2)

    return api_call_cost

# Get Video Thumbnail URL & Title
def video_info(url: str):
    """
    Get the thumbnail URL and title of a YouTube video.

    Parameters:
    url (str): The URL of the YouTube video.

    Returns:
    tuple: The thumbnail URL and title of the video.
    """

    yt = YouTube(url)

    # Get the thumbnail URL and title
    thumbnail_url = yt.thumbnail_url
    title = yt.title

    return thumbnail_url, title

# Download YouTube video as Audio
def download_audio(url: str):
    """
    Download the audio of a YouTube video.

    Parameters:
    url (str): The URL of the YouTube video.
    """
    yt = YouTube(url)

    # Extract the video_id from the url
    query = urlparse(url).query
    params = parse_qs(query)
    video_id = params["v"][0]

    # Get the first available audio stream and download it
    audio_stream = yt.streams.filter(only_audio=True).first()
    audio_stream.download(output_path="tmp/")

    # Convert the downloaded audio file to mp3 format
    audio_path = os.path.join("tmp/", audio_stream.default_filename)
    audio_clip = AudioFileClip(audio_path)
    audio_clip.write_audiofile(os.path.join("tmp/", f"{video_id}.mp3"))

    # Delete the original audio stream
    os.remove(audio_path)

# Transcription 
def transcribe_audio(file_path, video_id):
        """
        Transcribe the audio file of a YouTube video.

        Parameters:
        file_path (str): The path of the audio file.
        video_id (str): The ID of the YouTube video.
        """
        # The path of the transcript
        transcript_filepath = f"tmp/{video_id}.txt"

        # Get the size of the file in bytes
        file_size = os.path.getsize(file_path)

        # Convert bytes to megabytes
        file_size_in_mb = file_size / (1024 * 1024)

        # Check if the file size is less than 25 MB
        if file_size_in_mb < 25:
            print("Transcribing the audio file...(less than 25 mb) ")
            with open(file_path, "rb") as audio_file:
                transcript = openai.audio.transcriptions.create(model="whisper-1", file=audio_file, language="en")
                
                # Writing the content of transcript into a txt file
                with open(transcript_filepath, 'w') as transcript_file:
                    transcript_file.write(transcript.text)

            # Deleting the mp3 file
            os.remove(file_path)

        else:
            print("Transcribing the audio file...(more than 25 mb) ")
            # print("Please provide a smaller audio file (less than 25mb).")
            song = AudioSegment.from_mp3(file_path)

            # PyDub handles time in milliseconds
            minutes_chunk_24 = 24 * 60 * 1000

            # generate a unique random base chunk name for each session call based on datetime
            chunk_name = f"chunk_{datetime.now().strftime('%Y%m%d%H%M%S')}"
            chunk_no = 0
            for chunk in song[::minutes_chunk_24]:
                # chunk.export("good_morning_10.mp3", format="mp3")
                print("chunk : ",chunk)
                chunk_path = f"tmp/{chunk_name}_{chunk_no}.mp3"
                chunk.export(chunk_path, format="mp3")

                with open(chunk_path, "rb") as audio_file:
                    transcript = openai.audio.transcriptions.create(model="whisper-1", file=audio_file, language="en")
                    
                    # append the content of transcript into a txt file
                    with open(transcript_filepath, 'a') as transcript_file:
                        transcript_file.write(transcript.text)
                
                chunk_no += 1

            # delete the chunks after processing
            for i in range(chunk_no):
                os.remove(f"tmp/{chunk_name}_{i}.mp3")

def convert_youtube_link(youtube_link):
    """
    Convert a YouTube link to the standard format.

    Parameters:
    youtube_link (str): The YouTube link.

    Returns:
    str: The converted YouTube link.
    """
    # Check if the input is a valid YouTube link
    if "youtu.be" in youtube_link:
        # Extract the video ID from the input link
        video_id = youtube_link.split("/")[-1].split("?")[0]
        
        # Construct the new YouTube link format
        converted_link = f"https://www.youtube.com/watch?v={video_id}"
        
        return converted_link
    else:
        return youtube_link

@timeit  
def download_and_transcribe_audio(url: str):
    """
    Download and transcribe the audio of a YouTube video.

    Parameters:
    url (str): The URL of the YouTube video.

    Returns:
    str: The transcribed text of the audio.
    """
    # Extract the video_id from the url

    url = convert_youtube_link(url)
    query = urlparse(url).query
    params = parse_qs(query)
    video_id = params["v"][0]

    # The path of the audio file
    audio_path = f"tmp/{video_id}.mp3"

    # The path of the transcript
    transcript_filepath = f"tmp/{video_id}.txt"

    # Check if the transcript file already exist
    if os.path.exists(transcript_filepath):
        
        loader = TextLoader(transcript_filepath, encoding='utf8')
        documents = loader.load()

        time.sleep(5)

        print("*"*50, "DOCUMENTS: \n", documents, "*"*50)

        # return the text of the document
        return documents[0].page_content

    else: 
        print("downloading audio...")
        download_audio(url)
        print("audio downloaded...")
        # audio_path = "tmp/9JUAPgtkKpI.mp3"

        # Transcribe the mp3 audio to text
        transcribe_audio(audio_path, video_id)

        # Generating summary of the text file
        with open(transcript_filepath) as f:
            transcript_file = f.read()

        # return the text of the transcript file
        print("*"*50, "TRANSCRIPT: \n", transcript_file, "*"*50)
        return transcript_file


# Generate Answer
def generate_answer(api_key: str, url: str, question: str) -> str:
    """
    Generate an answer to a question based on the content of a YouTube video.

    Parameters:
    api_key (str): The OpenAI API key.
    url (str): The URL of the YouTube video.
    question (str): The question to answer.

    Returns:
    str: The generated answer.
    """
    openai.openai_api_key = api_key

    llm = OpenAI(temperature=0, model_name="gpt-3.5-turbo")
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)

    # Extract the video_id from the url
    query = urlparse(url).query
    params = parse_qs(query)
    video_id = params["v"][0]

    # The path of the audio file
    audio_path = f"tmp/{video_id}.mp3"

    # The path of the transcript
    transcript_filepath = f"tmp/{video_id}.txt"

    # Check if the transcript file already exist
    if os.path.exists(transcript_filepath):
        
        loader = TextLoader(transcript_filepath, encoding='utf8')
        documents = loader.load()

        print("*"*50, "DOCUMENTS: \n", documents, "*"*50)
        
        texts = text_splitter.split_documents(documents)
        embeddings = OpenAIEmbeddings()
        db = Chroma.from_documents(texts, embeddings)

        retriever = db.as_retriever()
        qa = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=retriever)
        answer = qa.run(question)

    else: 
        download_audio(url)

        # Transcribe the mp3 audio to text
        transcribe_audio(audio_path, video_id)

        # Generating summary of the text file
        loader = TextLoader(transcript_filepath, encoding='utf8')
        documents = loader.load()

        print("*"*50, "DOCUMENTS: \n", documents, "*"*50)
        
        texts = text_splitter.split_documents(documents)
        embeddings = OpenAIEmbeddings()
        db = Chroma.from_documents(texts, embeddings)

        retriever = db.as_retriever()
        qa = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=retriever)
        
        answer = qa.run(question)

    return answer.strip()
    

def generate_answer_from_text(text_path: str, question: str) -> str:
    """
    Generate an answer to a question based on the content of a text file.

    Parameters:
    text_path (str): The path of the text file.
    question (str): The question to answer.

    Returns:
    str: The generated answer.
    """
    llm = OpenAI(temperature=0, model_name="gpt-3.5-turbo")
    text_splitter = CharacterTextSplitter(chunk_size=100, chunk_overlap=0)

    loader = TextLoader(text_path, encoding='utf8')
    documents = loader.load()
    
    texts = text_splitter.split_documents(documents)
    embeddings = OpenAIEmbeddings()
    db = Chroma.from_documents(texts, embeddings)

    retriever = db.as_retriever()
    qa = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=retriever)
    print("%"*500,qa, "%"*500)
    answer = qa.run(question)

    return answer.strip()


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract the text from a PDF file.

    Parameters:
    pdf_path (str): The path of the PDF file.

    Returns:
    str: The extracted text.
    """
    import PyPDF2
    pdf = open(pdf_path, 'rb')
    pdfReader = PyPDF2.PdfReader(pdf)
    text_data = ""
    for i in range(len(pdfReader.pages)):
        page = pdfReader.pages[i]
        text = page.extract_text()
        text_data += " ".join(text.split("\t"))

    return text_data

def extract_text_from_doc(file_path: str) -> str:
    """
    Extract the text from a DOC file.

    Parameters:
    file_path (str): The path of the DOC file.

    Returns:
    str: The extracted text.
    """
    from docx import Document
    doc = Document(file_path)
    text_data = ""
    # Extract and print text from each paragraph
    for paragraph in doc.paragraphs:
        print(paragraph.text)
        text_data += " ".join(paragraph.text.split("\t"))
    print(f"############################ text: {text_data} #####################")
    return text_data



def generate_answer_from_text_anthropic(text_path: str, question: str) -> str:
    """
    Generate an answer to a question based on the content of a text file using the Anthropic API.

    Parameters:
    text_path (str): The path of the text file.
    question (str): The question to answer.

    Returns:
    str: The generated answer.
    """
    # anthropic.anthropic_api_key = api_key  # Set Anthropic API key
    ANTHROPIC_KEY = settings.ANTHROPIC_KEY
    
    # llm = anthropic.LLM(model_name="text-davinci-003")  # 
    llm = anthropic.Client(ANTHROPIC_KEY)
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)



    loader = TextLoader(text_path, encoding='utf8')
    documents = loader.load()
    
    texts = text_splitter.split_documents(documents)
    embeddings = OpenAIEmbeddings()
    db = Chroma.from_documents(texts, embeddings)

    retriever = db.as_retriever()
    qa = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=retriever)
    answer = qa.run(question)

    return answer.strip()
    


# Generating Video Summary 
def generate_summary(api_key: str, url: str) -> str:
    """
    Generate a summary of a YouTube video.

    Parameters:
    api_key (str): The OpenAI API key.
    url (str): The URL of the YouTube video.

    Returns:
    str: The generated summary.
    """
    openai.api_key = api_key

    llm = OpenAI(temperature=0, model_name="gpt-3.5-turbo")
    text_splitter = CharacterTextSplitter()

    # Extract the video_id from the url
    query = urlparse(url).query
    params = parse_qs(query)
    video_id = params["v"][0]

    # The path of the audio file
    audio_path = f"tmp/{video_id}.mp3"

    # The path of the transcript
    transcript_filepath = f"tmp/{video_id}.txt"

    # Check if the transcript file already exist
    if os.path.exists(transcript_filepath):
        # Generating summary of the text file
        with open(transcript_filepath) as f:
            transcript_file = f.read()

        texts = text_splitter.split_text(transcript_file)
        docs = [Document(page_content=t) for t in texts[:3]]
        chain = load_summarize_chain(llm, chain_type="map_reduce")
        summary = chain.run(docs)
    
    else: 
        download_audio(url)

        # Transcribe the mp3 audio to text
        transcribe_audio(audio_path, video_id)

        # Generating summary of the text file
        with open(transcript_filepath) as f:
            transcript_file = f.read()

        texts = text_splitter.split_text(transcript_file)
        docs = [Document(page_content=t) for t in texts[:3]]
        chain = load_summarize_chain(llm, chain_type="map_reduce")
        summary = chain.run(docs)

    return summary.strip()