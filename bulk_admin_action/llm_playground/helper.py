from datetime import datetime
import os
import shutil
from string import Template
import zipfile

from docx import Document

from commons.anthropic import anthropic_completion
from commons.google_apis import gemini_completion

def process_files(df, filename, llm_type='gemini'):
    project_root = os.path.dirname(os.path.abspath(__file__))
    output_folder = os.path.join(project_root,'..',"media")
    os.makedirs(output_folder, exist_ok=True)

    # Check and create 'Output' if not exists
    if 'Output' not in df.columns:
        df['Output'] = ""

    df['Output'] = df['Output'].astype('object')

    print("ABX")
    word_folder_path = os.path.join(output_folder, 'word_files')
    for index, row in df.iterrows():
        row_identifier = f"file-row-{index+1}"

        template = Template(row['Prompt'])
        prompt = template.safe_substitute(row)
        print(f"   ➤ Prompt: {prompt}")

        output = gemini_completion(prompt) if llm_type == 'gemini' else anthropic_completion(prompt, max_tokens=2024)
        df.at[index, 'Output'] = str(output)

        doc = Document()
        doc.add_heading(f'Row {index + 1}', level=1)

        for col in df.columns:
            doc.add_paragraph(col, style='Intense Quote')
            doc.add_paragraph(str(df.at[index, col]))

        safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in row_identifier[:50])
        doc_path = os.path.join(word_folder_path, f"{safe_title}.docx")

        # Ensure output folder still exists
        os.makedirs(word_folder_path, exist_ok=True)

        doc.save(doc_path)
        print(f"   📄 Saved Word: {safe_title}.docx")

    # Save CSV
    csv_filename = f"updated_{filename}.csv"
    df.to_csv(csv_filename, index=False)
    print(f"\n💾 CSV saved: {csv_filename}")

    # Create ZIP
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_filename = f"{project_root}/{filename[:10]}_bundle_{timestamp}.zip"
    with zipfile.ZipFile(zip_filename, 'w') as zipf:
        if csv_filename:
            zipf.write(csv_filename)
        for root, _, files_list in os.walk(word_folder_path):
            for file in files_list:
                file_path = os.path.join(root, file)
                zipf.write(file_path, arcname=file)
    print(f"📦 ZIP created: {zip_filename}")

    # Cleanup
    if csv_filename and os.path.exists(csv_filename):
        os.remove(csv_filename)
    if os.path.exists(output_folder):
        shutil.rmtree(output_folder)
    print("🧹 Cleanup complete. ", zip_filename)

    return zip_filename
